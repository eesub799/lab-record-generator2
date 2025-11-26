from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import qrcode
from PIL import Image
import os
from typing import List
from datetime import datetime
import io
import shutil

app = FastAPI(title="Lab Record Generator API")

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class Experiment(BaseModel):
    title: str
    date: str = ""
    github: str

class RecordData(BaseModel):
    course_title: str
    student_name: str
    register_number: str
    experiments: List[Experiment]

def create_qr_code(url: str, size: int = 200):
    """Generate QR code image"""
    qr = qrcode.QRCode(
        version=1,
        error_correction=qrcode.constants.ERROR_CORRECT_L,
        box_size=10,
        border=2,
    )
    qr.add_data(url)
    qr.make(fit=True)
    
    img = qr.make_image(fill_color="black", back_color="white")
    
    # Resize to specific size
    img = img.resize((size, size), Image.Resampling.LANCZOS)
    
    # Save to bytes
    img_byte_arr = io.BytesIO()
    img.save(img_byte_arr, format='PNG')
    img_byte_arr.seek(0)
    
    return img_byte_arr

def set_cell_border(cell, **kwargs):
    """Set cell border"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'tc{}'.format(edge.capitalize())
            element = OxmlElement('w:{}'.format(tag))
            element.set(qn('w:val'), 'single')
            element.set(qn('w:sz'), '4')
            element.set(qn('w:space'), '0')
            element.set(qn('w:color'), '000000')
            tcPr.append(element)

@app.get("/")
async def root():
    return {
        "message": "Lab Record Generator API",
        "status": "running",
        "version": "1.0",
        "logo_uploaded": os.path.exists("college_logo.png") or os.path.exists("college_logo.jpg") or os.path.exists("college_logo.jpeg")
    }

@app.post("/upload-logo")
async def upload_logo(file: UploadFile = File(...)):
    """Upload college logo"""
    try:
        # Validate file type
        if not file.content_type.startswith('image/'):
            raise HTTPException(status_code=400, detail="File must be an image")
        
        # Determine file extension
        ext = file.filename.split('.')[-1].lower()
        if ext not in ['png', 'jpg', 'jpeg']:
            raise HTTPException(status_code=400, detail="Only PNG, JPG, JPEG files allowed")
        
        # Save the logo
        logo_path = f"college_logo.{ext}"
        with open(logo_path, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        # Convert to PNG if needed (for consistency)
        if ext in ['jpg', 'jpeg']:
            img = Image.open(logo_path)
            img.save("college_logo.png")
            os.remove(logo_path)
            logo_path = "college_logo.png"
        
        return {
            "message": "Logo uploaded successfully",
            "filename": logo_path
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error uploading logo: {str(e)}")

@app.post("/generate-docx")
async def generate_docx(data: RecordData):
    try:
        # Create new document
        doc = Document()
        
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Add logo if exists
        logo_files = ['college_logo.png', 'college_logo.jpg', 'college_logo.jpeg']
        logo_exists = False
        for logo_file in logo_files:
            if os.path.exists(logo_file):
                logo_para = doc.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = logo_para.add_run()
                # Add logo with proper width (7.27 inches as per your sample)
                run.add_picture(logo_file, width=Inches(7.0))
                logo_exists = True
                break
        
        # Add spacing after logo
        if logo_exists:
            doc.add_paragraph()
        
        # Add course title (centered and bold)
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.add_run(data.course_title)
        title_run.bold = True
        title_run.font.size = Pt(14)
        
        # Add spacing
        doc.add_paragraph()
        
        # Create table
        num_experiments = len(data.experiments)
        table = doc.add_table(rows=num_experiments + 1, cols=6)
        table.style = 'Table Grid'
        
        # Header row
        headers = ['Exp', 'Date', 'Name of The Experiment', 'QR Code', 'Mark', 'Signature']
        header_cells = table.rows[0].cells
        
        for idx, header in enumerate(headers):
            cell = header_cells[idx]
            cell.text = header
            
            # Make header bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(11)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Set borders
            set_cell_border(cell, top={'val': 'single'}, bottom={'val': 'single'}, 
                          left={'val': 'single'}, right={'val': 'single'})
        
        # Set column widths
        widths = [Inches(0.5), Inches(0.8), Inches(3.0), Inches(0.8), Inches(0.6), Inches(1.0)]
        for idx, width in enumerate(widths):
            for row in table.rows:
                row.cells[idx].width = width
        
        # Fill data rows
        qr_images = []
        
        for idx, exp in enumerate(data.experiments):
            row = table.rows[idx + 1]
            cells = row.cells
            
            # Experiment number
            cells[0].text = str(idx + 1).zfill(2)
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Date
            cells[1].text = exp.date if exp.date else ""
            
            # Experiment title and GitHub link
            title_para = cells[2].paragraphs[0]
            title_para.text = exp.title
            title_para.add_run('\n\n')
            link_run = title_para.add_run(exp.github)
            link_run.font.size = Pt(9)
            link_run.font.color.rgb = RGBColor(0, 0, 255)
            link_run.underline = True
            
            # Generate and insert QR code
            qr_img_data = create_qr_code(exp.github, size=150)
            qr_filename = f"qr_{idx}.png"
            with open(qr_filename, 'wb') as f:
                f.write(qr_img_data.read())
            qr_images.append(qr_filename)
            
            # Clear cell and add image
            cells[3].text = ''
            paragraph = cells[3].paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(qr_filename, width=Inches(0.75))
            
            # Empty cells for Mark and Signature
            cells[4].text = ''
            cells[5].text = ''
            
            # Set borders for all cells
            for cell in cells:
                set_cell_border(cell, top={'val': 'single'}, bottom={'val': 'single'},
                              left={'val': 'single'}, right={'val': 'single'})
        
        # Add spacing after table
        doc.add_paragraph()
        doc.add_paragraph()
        
        # Add confirmation statement
        confirmation = doc.add_paragraph()
        conf_run = confirmation.add_run(
            'I confirm that the experiments and GitHub links provided are entirely my own work.'
        )
        conf_run.bold = True
        
        # Add spacing
        doc.add_paragraph()
        
        # Add student details in a table for proper alignment
        details_table = doc.add_table(rows=2, cols=2)
        details_table.autofit = False
        
        # First row: Name and Register Number
        name_cell = details_table.rows[0].cells[0]
        reg_cell = details_table.rows[0].cells[1]
        
        name_para = name_cell.paragraphs[0]
        name_para.add_run(f'Name: {data.student_name}').font.size = Pt(11)
        
        reg_para = reg_cell.paragraphs[0]
        reg_para.add_run(f'Register Number: {data.register_number}').font.size = Pt(11)
        reg_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Second row: Date and Learner's Signature
        date_cell = details_table.rows[1].cells[0]
        sign_cell = details_table.rows[1].cells[1]
        
        date_para = date_cell.paragraphs[0]
        date_para.add_run('Date:').font.size = Pt(11)
        
        sign_para = sign_cell.paragraphs[0]
        sign_para.add_run("Learner's Signature:").font.size = Pt(11)
        sign_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Remove table borders
        for row in details_table.rows:
            for cell in row.cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcBorders = OxmlElement('w:tcBorders')
                for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                    border = OxmlElement(f'w:{border_name}')
                    border.set(qn('w:val'), 'none')
                    tcBorders.append(border)
                tcPr.append(tcBorders)
        
        # Save document
        output_filename = f"{data.register_number}_Lab_Record.docx"
        doc.save(output_filename)
        
        # Return file
        response = FileResponse(
            path=output_filename,
            filename=output_filename,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
        # Cleanup QR images
        for qr_file in qr_images:
            try:
                if os.path.exists(qr_file):
                    os.remove(qr_file)
            except:
                pass
        
        return response
        
    except Exception as e:
        # Cleanup on error
        if 'qr_images' in locals():
            for qr_file in qr_images:
                try:
                    if os.path.exists(qr_file):
                        os.remove(qr_file)
                except:
                    pass
        
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 8000))
    uvicorn.run(app, host="0.0.0.0", port=port)